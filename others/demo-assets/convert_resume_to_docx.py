#!/usr/bin/env python3
"""Convert sample resume from markdown to DOCX format."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume_docx():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Priya Sharma', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    # Contact Information
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run(
        '📧 priya.sharma.demo@example.com | 📱 +91 98765 43210 | 📍 Pune, Maharashtra\n'
        'LinkedIn: linkedin.com/in/priya-sharma-demo'
    )
    contact_run.font.size = Pt(10)

    doc.add_paragraph()  # Spacing

    # Professional Summary
    doc.add_heading('Professional Summary', 1)
    summary = doc.add_paragraph(
        'Software developer with 2.5 years of experience in full-stack web development. '
        'Proficient in JavaScript, React, Node.js, and MongoDB. Built and deployed 5 production '
        'applications serving 10K+ users. Seeking to transition into UX design or product management '
        'to combine technical skills with user-centered problem solving.'
    )

    # Work Experience
    doc.add_heading('Work Experience', 1)

    # Job 1
    job1_title = doc.add_heading('Junior Software Developer', 2)
    job1_company = doc.add_paragraph()
    job1_company_run = job1_company.add_run('TechVista Solutions Pvt. Ltd., Pune | July 2022 - Present')
    job1_company_run.italic = True

    achievements1 = [
        'Developed and maintained 3 customer-facing web applications using React, Node.js, and PostgreSQL',
        'Collaborated with design team to implement responsive UI components, improving mobile user engagement by 35%',
        'Wrote automated tests using Jest and Cypress, reducing production bugs by 40%',
        'Participated in Agile sprints, daily standups, and bi-weekly retrospectives',
        'Mentored 2 junior interns on JavaScript fundamentals and Git workflows'
    ]

    for achievement in achievements1:
        p = doc.add_paragraph(achievement, style='List Bullet')

    doc.add_paragraph().add_run('Key Projects:').bold = True
    projects1 = [
        'E-commerce Dashboard: Built admin panel for inventory management using React and Material-UI',
        'Customer Portal: Developed authentication system with JWT and role-based access control',
        'Analytics Integration: Integrated Google Analytics and Mixpanel for user behavior tracking'
    ]
    for project in projects1:
        doc.add_paragraph(project, style='List Bullet 2')

    # Job 2
    job2_title = doc.add_heading('Software Development Intern', 2)
    job2_company = doc.add_paragraph()
    job2_company_run = job2_company.add_run('InnoSpark Technologies, Mumbai | January 2022 - June 2022')
    job2_company_run.italic = True

    achievements2 = [
        'Assisted in building REST APIs using Express.js and MongoDB',
        'Fixed 25+ bugs across front-end and back-end codebases',
        'Participated in code reviews and documentation updates',
        'Learned Docker basics and deployed test environments'
    ]

    for achievement in achievements2:
        doc.add_paragraph(achievement, style='List Bullet')

    # Education
    doc.add_heading('Education', 1)

    edu1_title = doc.add_heading('Bachelor of Engineering in Computer Science', 2)
    edu1_details = doc.add_paragraph()
    edu1_details_run = edu1_details.add_run('Savitribai Phule Pune University | 2018 - 2022')
    edu1_details_run.italic = True

    doc.add_paragraph('CGPA: 8.2/10', style='List Bullet')
    doc.add_paragraph(
        'Relevant Coursework: Data Structures, Web Technologies, Database Management, Human-Computer Interaction',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Final Year Project: "Smart Campus Navigation App" - Android application with real-time crowd-sourced location data',
        style='List Bullet'
    )

    edu2_title = doc.add_heading('Higher Secondary Certificate (HSC)', 2)
    edu2_details = doc.add_paragraph()
    edu2_details_run = edu2_details.add_run('Maharashtra State Board | 2018')
    edu2_details_run.italic = True

    doc.add_paragraph('Percentage: 87.6%', style='List Bullet')
    doc.add_paragraph('Stream: Science (Physics, Chemistry, Mathematics, Computer Science)', style='List Bullet')

    # Technical Skills
    doc.add_heading('Technical Skills', 1)

    skills = [
        ('Programming Languages', 'JavaScript (ES6+), Python, Java, SQL, HTML5, CSS3'),
        ('Frameworks & Libraries', 'React, Node.js, Express.js, Redux, Bootstrap, Tailwind CSS'),
        ('Databases', 'MongoDB, PostgreSQL, MySQL'),
        ('Tools & Platforms', 'Git, GitHub, VS Code, Postman, Docker (basic), Figma (basic), Jira, Slack'),
        ('Concepts', 'REST APIs, Agile/Scrum, Responsive Design, Authentication (JWT, OAuth), Testing (Jest, Cypress)')
    ]

    for category, items in skills:
        p = doc.add_paragraph()
        p.add_run(f'{category}: ').bold = True
        p.add_run(items)

    # Projects & Portfolio
    doc.add_heading('Projects & Portfolio', 1)

    project1_title = doc.add_heading('Personal Budget Tracker | 2023', 2)
    project1_details = [
        'Built a full-stack web app for tracking income and expenses with data visualization',
        'Tech Stack: React, Node.js, MongoDB, Chart.js',
        'Deployed on Heroku with CI/CD pipeline',
        'Link: github.com/priya-sharma-demo/budget-tracker'
    ]
    for detail in project1_details:
        doc.add_paragraph(detail, style='List Bullet')

    project2_title = doc.add_heading('Recipe Sharing Platform | 2022', 2)
    project2_details = [
        'Community-driven recipe app with user ratings and comments',
        'Tech Stack: React, Firebase, Material-UI',
        'Implemented real-time notifications and image uploads',
        'Link: github.com/priya-sharma-demo/recipe-hub'
    ]
    for detail in project2_details:
        doc.add_paragraph(detail, style='List Bullet')

    # Certifications
    doc.add_heading('Certifications', 1)
    certs = [
        'Meta Front-End Developer Professional Certificate | Coursera | 2023',
        'AWS Cloud Practitioner Essentials | AWS Training | 2023',
        'JavaScript Algorithms and Data Structures | freeCodeCamp | 2021'
    ]
    for cert in certs:
        doc.add_paragraph(cert, style='List Bullet')

    # Soft Skills & Interests
    doc.add_heading('Soft Skills & Interests', 1)

    doc.add_paragraph().add_run('Strengths:').bold = True
    strengths = [
        'Strong problem-solving and analytical thinking',
        'Good communication and teamwork',
        'Quick learner with curiosity for emerging technologies',
        'Attention to detail in code quality and user experience'
    ]
    for strength in strengths:
        doc.add_paragraph(strength, style='List Bullet')

    doc.add_paragraph().add_run('Interests:').bold = True
    interests = [
        'UI/UX design and user research',
        'Product strategy and roadmapping',
        'Reading tech blogs and design case studies',
        'Attending tech meetups and hackathons (won 2nd place at PuneHacks 2023)'
    ]
    for interest in interests:
        doc.add_paragraph(interest, style='List Bullet')

    doc.add_paragraph().add_run('Languages:').bold = True
    languages = [
        'English (Fluent)',
        'Hindi (Native)',
        'Marathi (Native)'
    ]
    for lang in languages:
        doc.add_paragraph(lang, style='List Bullet')

    # Career Aspirations
    doc.add_heading('Career Aspirations', 1)

    aspirations_intro = doc.add_paragraph(
        "I'm excited about transitioning into roles that bridge technology and user needs. "
        "My ideal next step would be:"
    )

    aspirations = [
        'UX/UI Designer - Leverage my technical background to design intuitive, user-centered interfaces and collaborate closely with development teams',
        'Product Manager - Combine technical expertise with strategic thinking to define product roadmaps and drive user impact',
        'Front-End Lead - Deepen front-end specialization with focus on accessibility, performance, and design systems'
    ]

    for i, aspiration in enumerate(aspirations, 1):
        doc.add_paragraph(f'{i}. {aspiration}')

    doc.add_paragraph(
        'I value work-life balance, continuous learning opportunities, and working on products that make a meaningful difference. '
        'Open to remote or hybrid roles in Pune, Mumbai, or Bangalore.'
    )

    # References
    doc.add_heading('References', 1)
    doc.add_paragraph('Available upon request')

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('Resume Last Updated: August 2026')
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True

    # Save document
    output_path = 'Priya_Sharma_Resume_Demo.docx'
    doc.save(output_path)
    print(f'✅ Resume saved as {output_path}')
    return output_path

if __name__ == '__main__':
    create_resume_docx()
